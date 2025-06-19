#!/usr/bin/env python3
"""
Script de conversion Markdown vers Word (DOCX)
Convertit le plan de maintenance Cesizen du format MD vers DOCX
"""

import os
import sys
import re
from pathlib import Path

try:
    import markdown
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.shared import OxmlElement, qn
except ImportError as e:
    print(f"Erreur : Module manquant - {e}")
    print("Installation nécessaire :")
    print("pip install python-docx markdown")
    sys.exit(1)

class MarkdownToDocxConverter:
    def __init__(self):
        self.doc = Document()
        self.setup_styles()
    
    def setup_styles(self):
        """Configure les styles personnalisés pour le document"""
        
        # Style pour le titre principal
        title_style = self.doc.styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
        title_font = title_style.font
        title_font.name = 'Arial'
        title_font.size = Pt(24)
        title_font.bold = True
        title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_style.paragraph_format.space_after = Pt(24)
        
        # Style pour les titres de niveau 2
        h2_style = self.doc.styles.add_style('CustomH2', WD_STYLE_TYPE.PARAGRAPH)
        h2_font = h2_style.font
        h2_font.name = 'Arial'
        h2_font.size = Pt(16)
        h2_font.bold = True
        h2_style.paragraph_format.space_before = Pt(16)
        h2_style.paragraph_format.space_after = Pt(12)
        
        # Style pour les titres de niveau 3
        h3_style = self.doc.styles.add_style('CustomH3', WD_STYLE_TYPE.PARAGRAPH)
        h3_font = h3_style.font
        h3_font.name = 'Arial'
        h3_font.size = Pt(14)
        h3_font.bold = True
        h3_style.paragraph_format.space_before = Pt(12)
        h3_style.paragraph_format.space_after = Pt(8)
        
        # Style pour le texte normal
        normal_style = self.doc.styles['Normal']
        normal_font = normal_style.font
        normal_font.name = 'Arial'
        normal_font.size = Pt(11)
        normal_style.paragraph_format.space_after = Pt(6)
        normal_style.paragraph_format.line_spacing = 1.15
    
    def add_header_footer(self):
        """Ajoute en-tête et pied de page"""
        # En-tête
        header = self.doc.sections[0].header
        header_para = header.paragraphs[0]
        header_para.text = "Plan de Maintenance - Projet Cesizen"
        header_para.style = self.doc.styles['Header']
        
        # Pied de page avec numérotation
        footer = self.doc.sections[0].footer
        footer_para = footer.paragraphs[0]
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_para.text = "Page "
        
        # Ajout du numéro de page
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        footer_para._element.append(fldChar1)
        
        instrText = OxmlElement('w:instrText')
        instrText.text = "PAGE"
        footer_para._element.append(instrText)
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        footer_para._element.append(fldChar2)
    
    def parse_markdown_line(self, line):
        """Parse une ligne markdown et retourne le type et le contenu"""
        line = line.strip()
        
        if not line:
            return 'empty', ''
        
        # Titre principal
        if line.startswith('# '):
            return 'h1', line[2:]
        
        # Titre niveau 2
        if line.startswith('## '):
            return 'h2', line[3:]
        
        # Titre niveau 3
        if line.startswith('### '):
            return 'h3', line[4:]
        
        # Titre niveau 4
        if line.startswith('#### '):
            return 'h4', line[5:]
        
        # Liste à puces
        if line.startswith('- ') or line.startswith('* '):
            return 'bullet', line[2:]
        
        # Liste numérotée
        if re.match(r'^\d+\. ', line):
            return 'numbered', re.sub(r'^\d+\. ', '', line)
        
        # Tableau
        if '|' in line and line.count('|') >= 2:
            return 'table', line
        
        # Code block
        if line.startswith('```'):
            return 'code_block', line[3:]
        
        # Ligne horizontale
        if line.startswith('---'):
            return 'hr', ''
        
        # Texte normal
        return 'text', line
    
    def add_table_of_contents(self, lines):
        """Génère une table des matières basée sur les titres"""
        toc_paragraph = self.doc.add_paragraph()
        toc_paragraph.style = self.doc.styles['CustomH2']
        toc_paragraph.add_run("Table des matières").bold = True
        
        toc_items = []
        for line in lines:
            line_type, content = self.parse_markdown_line(line)
            if line_type in ['h2', 'h3']:
                level = 1 if line_type == 'h2' else 2
                toc_items.append((level, content))
        
        for level, title in toc_items:
            toc_para = self.doc.add_paragraph()
            if level == 1:
                toc_para.add_run(f"• {title}")
            else:
                toc_para.add_run(f"  ◦ {title}")
            toc_para.paragraph_format.left_indent = Inches(0.25 * level)
        
        self.doc.add_page_break()
    
    def process_table(self, table_lines):
        """Traite les lignes d'un tableau markdown"""
        if len(table_lines) < 2:
            return
        
        # Parse les en-têtes
        headers = [cell.strip() for cell in table_lines[0].split('|')[1:-1]]
        
        # Skip la ligne de séparation
        data_lines = table_lines[2:]
        
        # Crée le tableau
        table = self.doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'        # Ajoute les en-têtes avec formatage en gras forcé
        header_row = table.rows[0]
        for i, header in enumerate(headers):
            cell_para = header_row.cells[i].paragraphs[0]
            cell_para.clear()  # Nettoyer le contenu existant
            
            # Créer un run en gras pour l'en-tête
            header_run = cell_para.add_run(header.replace('**', ''))  # Supprimer les ** s'il y en a
            header_run.bold = True
            header_run.font.bold = True
        
        # Ajoute les données avec formatage
        for line in data_lines:
            if '|' in line:
                cells = [cell.strip() for cell in line.split('|')[1:-1]]
                if len(cells) == len(headers):
                    row = table.add_row()
                    for i, cell_content in enumerate(cells):
                        cell_para = row.cells[i].paragraphs[0]
                        cell_para.clear()  # Nettoyer le contenu existant                        self.add_formatted_run(cell_para, cell_content)
    
    def add_formatted_run(self, paragraph, text):
        """Ajoute un run formaté à un paragraphe en traitant le markdown (gras)"""
        if '**' not in text:
            # Pas de formatage, ajouter directement
            paragraph.add_run(text)
            return
        
        # Diviser le texte par les marqueurs **
        parts = text.split('**')
        
        for i, part in enumerate(parts):
            if not part:  # Ignorer les parties vides
                continue
                
            run = paragraph.add_run(part)
            
            # Si c'est un index impair, c'est du texte entre **
            if i % 2 == 1:
                run.bold = True
                run.font.bold = True  # Double application pour compatibilité

    def convert_file(self, input_file, output_file):
        """Convertit un fichier markdown en docx"""
        
        print(f"Lecture du fichier : {input_file}")
        
        try:
            with open(input_file, 'r', encoding='utf-8') as f:
                lines = f.readlines()
        except FileNotFoundError:
            print(f"Erreur : Fichier {input_file} non trouvé")
            return False
        except Exception as e:
            print(f"Erreur lors de la lecture : {e}")
            return False
          # Configuration du document
        self.add_header_footer()
        
        # Variables de traitement
        in_code_block = False
        in_table = False
        table_lines = []
        
        for line in lines:
            line_type, content = self.parse_markdown_line(line)
              if line_type == 'empty':
                if not in_code_block and not in_table:
                    continue
            
            elif line_type == 'h1':
                # Titre principal
                para = self.doc.add_paragraph()
                para.style = self.doc.styles['CustomTitle']
                title_run = para.add_run(content)
                title_run.bold = True
                title_run.font.bold = True
            
            elif line_type == 'h2':
                if in_table:
                    self.process_table(table_lines)
                    in_table = False
                    table_lines = []
                
                para = self.doc.add_paragraph()
                para.style = self.doc.styles['CustomH2']
                h2_run = para.add_run(content)
                h2_run.bold = True
                h2_run.font.bold = True
            
            elif line_type == 'h3':
                if in_table:
                    self.process_table(table_lines)
                    in_table = False
                    table_lines = []
                
                para = self.doc.add_paragraph()
                para.style = self.doc.styles['CustomH3']
                h3_run = para.add_run(content)
                h3_run.bold = True
                h3_run.font.bold = True
            
            elif line_type == 'h4':
                if in_table:
                    self.process_table(table_lines)
                    in_table = False
                    table_lines = []
                
                para = self.doc.add_paragraph(content)
                para.runs[0].bold = True
            
            elif line_type == 'bullet':
                if in_table:
                    self.process_table(table_lines)
                    in_table = False
                    table_lines = []
                
                para = self.doc.add_paragraph(content, style='List Bullet')
            
            elif line_type == 'numbered':
                if in_table:
                    self.process_table(table_lines)
                    in_table = False
                    table_lines = []
                
                para = self.doc.add_paragraph(content, style='List Number')
            
            elif line_type == 'table':
                if not in_table:
                    in_table = True
                    table_lines = []
                table_lines.append(line.strip())
            
            elif line_type == 'code_block':
                if content.strip():  # Code block avec langage
                    in_code_block = not in_code_block
                    if not in_code_block:
                        self.doc.add_paragraph()  # Espace après le code
                else:  # Code block vide
                    in_code_block = not in_code_block
            
            elif line_type == 'hr':
                if in_table:
                    self.process_table(table_lines)
                    in_table = False
                    table_lines = []
                
                # Ajouter une ligne horizontale (saut de page)
                self.doc.add_page_break()
            
            elif line_type == 'text':
                if in_table:
                    # Les lignes de texte dans un tableau sont ignorées
                    continue
                
                if in_code_block:
                    # Texte dans un bloc de code                    para = self.doc.add_paragraph(content)
                    para.style = 'No Spacing'
                    for run in para.runs:
                        run.font.name = 'Courier New'
                        run.font.size = Pt(10)
                else:
                    # Texte normal avec formatage Markdown
                    para = self.doc.add_paragraph()
                    self.add_formatted_run(para, content)
        
        # Traiter le dernier tableau si nécessaire
        if in_table and table_lines:
            self.process_table(table_lines)
        
        # Sauvegarde
        try:
            self.doc.save(output_file)
            print(f"Conversion réussie : {output_file}")
            return True
        except Exception as e:
            print(f"Erreur lors de la sauvegarde : {e}")
            return False

def main():
    """Fonction principale"""
    
    # Fichiers d'entrée et de sortie
    input_file = "plan_maintenance_cesizen.md"
    output_file = "Plan_Maintenance_Cesizen.docx"
    
    # Vérification de l'existence du fichier d'entrée
    if not os.path.exists(input_file):
        print(f"Erreur : Le fichier {input_file} n'existe pas.")
        print("Assurez-vous que le script est dans le même répertoire que le fichier markdown.")
        return
    
    # Conversion
    converter = MarkdownToDocxConverter()
    success = converter.convert_file(input_file, output_file)
    
    if success:
        print(f"\n✅ Conversion terminée avec succès !")
        print(f"📄 Fichier créé : {output_file}")
        print(f"📂 Emplacement : {os.path.abspath(output_file)}")
    else:
        print("\n❌ Échec de la conversion")

if __name__ == "__main__":
    main()
